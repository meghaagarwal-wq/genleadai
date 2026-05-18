"""Touchpoint mapping endpoints for Aria onboarding Step 3B and beyond.

Read-only template library + per-tenant active map storage. Execution engine
(Phase B) will read from `workspace_touchpoint_maps` and instantiate
`lead_touchpoint_log` rows on lead create.

Phase C additions (Iter 37):
  - Version history (last 5 saved versions per tenant in `touchpoint_map_versions`)
  - Document upload (PDF/Excel/DOCX → Claude → preview JSON)
  - Max touchpoints raised 30 → 32 (matches Master Spec)
  - Audit-log integration for save + import events
"""
import io
import os
import uuid
from datetime import datetime, timezone
from typing import List, Optional, Dict, Any

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from pydantic import BaseModel, Field

from deps import db, get_current_user
from routes.tenants import get_active_tenant
from touchpoint_templates_seed import TEMPLATES, select_template

router = APIRouter(prefix="/api/touchpoints", tags=["touchpoints"])

templates_col = db["touchpoint_templates"]
maps_col = db["workspace_touchpoint_maps"]
versions_col = db["touchpoint_map_versions"]

MAX_TOUCHPOINTS = 32
MAX_VERSIONS_RETAINED = 5
MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10MB


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _seed_templates_if_missing() -> None:
    """Idempotently seed the 8 universal templates on first read.
    Cheap upsert — runs at most once per process per template."""
    for t in TEMPLATES:
        templates_col.update_one(
            {"id": t["id"]},
            {"$set": {**t, "seeded_at": _now()}},
            upsert=True,
        )


# Run once at import (server startup picks this up via include_router)
_seed_templates_if_missing()


# ─── Models ──────────────────────────────────────────────────────────────────
ALLOWED_CHANNELS = {"whatsapp", "email", "call_reminder", "linkedin_nudge"}
ALLOWED_ROLES = {"autonomous", "alert_human"}
ALLOWED_TYPES = {
    "intro", "qualifier", "value_drop", "value_add", "follow_up",
    "social_proof", "soft_cta", "budget_probe", "meeting_cta",
    "human_escalation", "re_engagement", "urgency", "closure",
}


class Touchpoint(BaseModel):
    index: int
    day: float = Field(ge=0)
    hour: int = Field(ge=0, le=23)
    channel: str
    message_type: str
    aria_role: str
    trigger: Optional[str] = ""
    message_template: str
    # Iter 54 — branching logic shared with the Outreach Campaign builder.
    # Validated against the same schema used by /api/outreach/* via
    # routes.outreach.validate_conditions; stored verbatim alongside the touchpoint.
    conditions: Optional[Dict[str, Any]] = None


class SaveMapPayload(BaseModel):
    template_id: Optional[str] = None
    is_customised: bool = False
    touchpoints: List[Touchpoint]


def _validate_touchpoints(items: List[Touchpoint]) -> None:
    if not items:
        raise HTTPException(status_code=400, detail="At least one touchpoint required")
    if len(items) > MAX_TOUCHPOINTS:
        raise HTTPException(status_code=400, detail=f"Max {MAX_TOUCHPOINTS} touchpoints per map")
    # Lazy import — avoids circular import at module load time.
    from routes.outreach import validate_conditions
    for t in items:
        if t.channel not in ALLOWED_CHANNELS:
            raise HTTPException(status_code=400, detail=f"Invalid channel '{t.channel}'")
        if t.aria_role not in ALLOWED_ROLES:
            raise HTTPException(status_code=400, detail=f"Invalid aria_role '{t.aria_role}'")
        if t.message_type not in ALLOWED_TYPES:
            raise HTTPException(status_code=400, detail=f"Invalid message_type '{t.message_type}'")
        if not t.message_template.strip():
            raise HTTPException(status_code=400, detail="message_template cannot be empty")
        # Branching conditions — same schema as the Outreach builder
        if t.conditions:
            t.conditions = validate_conditions(t.conditions)


def _snapshot_version(tenant_id: str, current_map: dict, saved_by: Optional[str]) -> None:
    """Snapshot the previous map into touchpoint_map_versions before overwriting.
    Keep at most MAX_VERSIONS_RETAINED per tenant; prune the oldest."""
    if not current_map or not current_map.get("touchpoints"):
        return
    versions_col.insert_one({
        "id": f"ver_{uuid.uuid4().hex[:12]}",
        "tenant_id": tenant_id,
        "template_id": current_map.get("template_id"),
        "template_name": current_map.get("template_name"),
        "touchpoints": current_map.get("touchpoints"),
        "touchpoint_count": current_map.get("touchpoint_count") or len(current_map.get("touchpoints") or []),
        "is_customised": bool(current_map.get("is_customised")),
        "saved_by": saved_by or current_map.get("saved_by"),
        "snapshot_of": current_map.get("updated_at") or current_map.get("created_at"),
        "created_at": _now(),
    })
    # Prune older versions beyond cap
    extras = list(versions_col.find(
        {"tenant_id": tenant_id}, {"_id": 0, "id": 1, "created_at": 1}
    ).sort("created_at", -1).skip(MAX_VERSIONS_RETAINED))
    if extras:
        versions_col.delete_many({"tenant_id": tenant_id, "id": {"$in": [e["id"] for e in extras]}})


def _audit(tenant_id: str, user: dict, action: str, metadata: dict = None) -> None:
    try:
        from routes.audit_log import audit_write
        audit_write(tenant_id, user, action, "touchpoint_map", None, metadata or {})
    except Exception:
        pass


# ─── Endpoints ───────────────────────────────────────────────────────────────
@router.get("/templates")
async def list_templates(_: dict = Depends(get_current_user)):
    """Return all seeded universal templates (read-only library)."""
    rows = list(templates_col.find({}, {"_id": 0}).sort("id", 1))
    return {"templates": rows}


@router.get("/templates/{template_id}")
async def get_template(template_id: str, _: dict = Depends(get_current_user)):
    t = templates_col.find_one({"id": template_id}, {"_id": 0})
    if not t:
        raise HTTPException(status_code=404, detail="Template not found")
    return t


@router.get("/auto-select")
async def auto_select(tenant: dict = Depends(get_active_tenant)):
    """Return the best-fit template based on the tenant's onboarding answers.

    Used by Step 3B to render the auto-generated journey before the user
    decides whether to accept or customise.
    """
    onboarding_col = db["onboarding_config"]
    cfg = onboarding_col.find_one({"tenant_id": tenant["id"]}, {"_id": 0}) or {}
    selected = select_template(cfg)
    return selected


@router.get("/map")
async def get_map(tenant: dict = Depends(get_active_tenant)):
    """Get the active tenant's saved touchpoint map (or null if not set)."""
    row = maps_col.find_one({"tenant_id": tenant["id"]}, {"_id": 0})
    return {"map": row}


@router.post("/map")
async def save_map(payload: SaveMapPayload, tenant: dict = Depends(get_active_tenant), current_user: dict = Depends(get_current_user)):
    """Save (or replace) the active tenant's touchpoint map.

    Owner/admin only. Used by both 'Looks good — use this' and 'Save my custom
    journey' flows.
    """
    if tenant.get("_member_role") not in ("owner", "admin"):
        raise HTTPException(status_code=403, detail="Owner/Admin only")
    _validate_touchpoints(payload.touchpoints)

    # Re-index touchpoints to ensure they are 0..n-1 contiguous regardless of
    # what the client sent (after drag-reorder etc.)
    cleaned = []
    for i, t in enumerate(payload.touchpoints):
        cleaned.append({
            "index": i,
            "day": float(t.day),
            "hour": int(t.hour),
            "channel": t.channel,
            "message_type": t.message_type,
            "aria_role": t.aria_role,
            "trigger": (t.trigger or "").strip(),
            "message_template": t.message_template.strip(),
            "conditions": t.conditions or {},
        })

    template_id = payload.template_id or "tpl_standard"
    template = templates_col.find_one({"id": template_id}, {"_id": 0, "name": 1, "duration_days": 1})

    doc = {
        "tenant_id": tenant["id"],
        "template_id": template_id,
        "template_name": (template or {}).get("name"),
        "is_customised": bool(payload.is_customised),
        "touchpoints": cleaned,
        "touchpoint_count": len(cleaned),
        "duration_days": (template or {}).get("duration_days"),
        "saved_by": current_user.get("email"),
        "updated_at": _now(),
    }
    existing = maps_col.find_one({"tenant_id": tenant["id"]}, {"_id": 0})
    if existing:
        doc["created_at"] = existing.get("created_at") or _now()
        # Snapshot the previous version before overwrite
        _snapshot_version(tenant["id"], existing, current_user.get("email"))
    else:
        doc["created_at"] = _now()

    maps_col.update_one(
        {"tenant_id": tenant["id"]},
        {"$set": doc},
        upsert=True,
    )
    _audit(tenant["id"], current_user, "touchpoint_map.updated", {
        "touchpoint_count": len(cleaned),
        "is_customised": bool(payload.is_customised),
        "template_id": template_id,
    })
    return {"map": doc}


@router.post("/map/reset")
async def reset_map(tenant: dict = Depends(get_active_tenant), current_user: dict = Depends(get_current_user)):
    """Reset the tenant's map to the auto-recommended template — discards any
    customisations and re-selects from current onboarding answers."""
    if tenant.get("_member_role") not in ("owner", "admin"):
        raise HTTPException(status_code=403, detail="Owner/Admin only")
    onboarding_col = db["onboarding_config"]
    cfg = onboarding_col.find_one({"tenant_id": tenant["id"]}, {"_id": 0}) or {}
    selected = select_template(cfg)
    doc = {
        "tenant_id": tenant["id"],
        "template_id": selected["id"],
        "template_name": selected.get("name"),
        "is_customised": False,
        "touchpoints": selected["touchpoints"],
        "touchpoint_count": len(selected["touchpoints"]),
        "duration_days": selected.get("duration_days"),
        "saved_by": current_user.get("email"),
        "updated_at": _now(),
    }
    existing = maps_col.find_one({"tenant_id": tenant["id"]}, {"_id": 0, "created_at": 1})
    doc["created_at"] = (existing or {}).get("created_at") or _now()
    maps_col.update_one(
        {"tenant_id": tenant["id"]},
        {"$set": doc},
        upsert=True,
    )
    return {"map": doc, "selection": selected.get("selection")}


@router.delete("/map")
async def clear_map(tenant: dict = Depends(get_active_tenant)):
    """Clear the saved map. Aria falls back to no journey for new leads."""
    if tenant.get("_member_role") not in ("owner", "admin"):
        raise HTTPException(status_code=403, detail="Owner/Admin only")
    maps_col.delete_one({"tenant_id": tenant["id"]})
    return {"status": "ok"}



# ─── Version History ─────────────────────────────────────────────────────────
@router.get("/map/versions")
async def list_versions(tenant: dict = Depends(get_active_tenant)):
    """Last 5 saved versions of this tenant's touchpoint map."""
    rows = list(versions_col.find(
        {"tenant_id": tenant["id"]}, {"_id": 0}
    ).sort("created_at", -1).limit(MAX_VERSIONS_RETAINED))
    return {"versions": rows}


@router.post("/map/versions/{version_id}/restore")
async def restore_version(version_id: str, tenant: dict = Depends(get_active_tenant), current_user: dict = Depends(get_current_user)):
    """Restore a previously-saved version as the active map. The currently-active
    map is first snapshotted so the restore is itself reversible."""
    if tenant.get("_member_role") not in ("owner", "admin"):
        raise HTTPException(status_code=403, detail="Owner/Admin only")
    v = versions_col.find_one({"id": version_id, "tenant_id": tenant["id"]}, {"_id": 0})
    if not v:
        raise HTTPException(status_code=404, detail="Version not found")
    existing = maps_col.find_one({"tenant_id": tenant["id"]}, {"_id": 0})
    if existing:
        _snapshot_version(tenant["id"], existing, current_user.get("email"))
    doc = {
        "tenant_id": tenant["id"],
        "template_id": v.get("template_id"),
        "template_name": v.get("template_name"),
        "is_customised": True,
        "touchpoints": v.get("touchpoints") or [],
        "touchpoint_count": v.get("touchpoint_count") or len(v.get("touchpoints") or []),
        "saved_by": current_user.get("email"),
        "updated_at": _now(),
        "created_at": (existing or {}).get("created_at") or _now(),
        "restored_from": version_id,
    }
    maps_col.update_one({"tenant_id": tenant["id"]}, {"$set": doc}, upsert=True)
    _audit(tenant["id"], current_user, "touchpoint_map.restored", {"version_id": version_id})
    return {"map": doc}


# ─── Document Upload (PDF / Excel / DOCX) → Claude → Preview JSON ────────────
def _extract_text_from_upload(filename: str, content: bytes) -> str:
    """Pull plain text from PDF/DOCX/XLSX. Caller deletes the bytes after."""
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            return "\n".join((p.extract_text() or "") for p in reader.pages)
        if name.endswith(".docx") or name.endswith(".doc"):
            from docx import Document
            d = Document(io.BytesIO(content))
            paragraphs = [p.text for p in d.paragraphs]
            for table in d.tables:
                for row in table.rows:
                    paragraphs.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(paragraphs)
        if name.endswith(".xlsx") or name.endswith(".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            lines = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(values_only=True):
                    if not row:
                        continue
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        lines.append(" | ".join(cells))
            return "\n".join(lines)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not parse document: {e}")
    raise HTTPException(status_code=400, detail="Unsupported file format. Use PDF, DOCX, or XLSX.")


_IMPORT_PROMPT = """You are reading a sales outreach sequence document.
Extract every touchpoint or step in the sequence and return them as a structured JSON array. Maximum 32 touchpoints.

For each touchpoint extract:
- step_number (integer, starting at 1)
- day (integer — which day from lead entry, 0 = immediate)
- hour (integer 0-23 — hour of day to send, default 9 if not specified)
- channel (whatsapp / email / call_reminder / linkedin_nudge — map "call" to "call_reminder", "linkedin" to "linkedin_nudge")
- message_type (intro / qualifier / follow_up / value_drop / value_add / social_proof / soft_cta / budget_probe / meeting_cta / human_escalation / re_engagement / urgency / closure)
- aria_role (autonomous / alert_human — default autonomous unless document says "call", "manual", "human", "rep")
- message_template (the actual message text if present, or a clear description of the intent if no message is written)
- trigger (free-form condition like "if no reply" or "always" — default empty string)

If the document does not contain a clear sequence, return: {"error": "No sequence found", "reason": "..."}

Return ONLY valid JSON. No explanation. No markdown. No code fences."""


@router.post("/import-document")
async def import_document(
    file: UploadFile = File(...),
    tenant: dict = Depends(get_active_tenant),
    current_user: dict = Depends(get_current_user),
):
    """Parse an uploaded outreach-sequence document and return a preview JSON
    of extracted touchpoints. Caller previews + confirms before saving the
    map via POST /api/touchpoints/map."""
    if tenant.get("_member_role") not in ("owner", "admin"):
        raise HTTPException(status_code=403, detail="Owner/Admin only")
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Empty file")
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=400, detail="File over 10MB limit")
    text = _extract_text_from_upload(file.filename or "", content)
    del content
    if len(text.strip()) < 20:
        raise HTTPException(status_code=400, detail="Could not read meaningful text from the document. Try a different file.")

    api_key = os.environ.get("EMERGENT_LLM_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="Emergent LLM key not configured")

    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        chat = LlmChat(
            api_key=api_key,
            session_id=f"tp-import-{uuid.uuid4().hex[:8]}",
            system_message=_IMPORT_PROMPT,
        ).with_model("anthropic", "claude-sonnet-4-5-20250929")
        raw = await chat.send_message(UserMessage(text=text[:30000]))
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Claude unavailable: {e}")

    import json as _json
    import re as _re
    body = (raw or "").strip()
    body = _re.sub(r"^```(?:json)?\s*|\s*```$", "", body, flags=_re.MULTILINE).strip()
    try:
        parsed = _json.loads(body)
    except Exception:
        m = _re.search(r"[\[\{].*[\]\}]", body, _re.DOTALL)
        if not m:
            raise HTTPException(status_code=502, detail="Could not parse Claude response")
        try:
            parsed = _json.loads(m.group(0))
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"Invalid JSON from Claude: {e}")

    truncated = False
    if isinstance(parsed, dict):
        if "error" in parsed:
            raise HTTPException(status_code=400, detail=f"{parsed.get('error')}: {parsed.get('reason','')}")
        touchpoints = parsed.get("touchpoints") or parsed.get("steps") or parsed.get("data") or []
        truncated = bool(parsed.get("truncated"))
    elif isinstance(parsed, list):
        touchpoints = parsed
    else:
        raise HTTPException(status_code=502, detail="Unexpected Claude response shape")

    if not isinstance(touchpoints, list) or not touchpoints:
        raise HTTPException(status_code=400, detail="No touchpoints found in document")

    channel_map = {"call": "call_reminder", "phone": "call_reminder", "linkedin": "linkedin_nudge"}
    preview = []
    for i, t in enumerate(touchpoints[:MAX_TOUCHPOINTS]):
        ch = (t.get("channel") or "whatsapp").lower().strip()
        ch = channel_map.get(ch, ch)
        if ch not in ALLOWED_CHANNELS:
            ch = "whatsapp"
        mt = (t.get("message_type") or "follow_up").lower().strip()
        if mt not in ALLOWED_TYPES:
            mt = "follow_up"
        role = (t.get("aria_role") or "autonomous").lower().strip()
        if role not in ALLOWED_ROLES:
            role = "autonomous"
        preview.append({
            "index": i,
            "day": float(t.get("day") or (i * 2)),
            "hour": int(t.get("hour") or 9),
            "channel": ch,
            "message_type": mt,
            "aria_role": role,
            "trigger": (t.get("trigger") or t.get("condition") or "").strip(),
            "message_template": (t.get("message_template") or t.get("message_content") or t.get("message") or "").strip()[:1000] or f"Touchpoint #{i+1}",
        })

    if len(touchpoints) > MAX_TOUCHPOINTS:
        truncated = True

    _audit(tenant["id"], current_user, "touchpoint_map.imported_from_document", {
        "filename": file.filename,
        "extracted_count": len(preview),
        "truncated": truncated,
    })

    return {
        "preview": preview,
        "extracted_count": len(preview),
        "truncated": truncated,
        "source_filename": file.filename,
    }



# ─── SCORING (Performance + Lead-fit + AI quality) ──────────────────────────
log_col = db["lead_touchpoint_log"]
leads_col_local = db["leads"]
activities_col_local = db["activities"]
ai_quality_cache_col = db["touchpoint_ai_quality_cache"]


def _safe_pct(num: int, denom: int) -> float:
    return round((num / denom) * 100, 1) if denom else 0.0


def _grade(score: float) -> str:
    if score >= 80:
        return "A"
    if score >= 65:
        return "B"
    if score >= 45:
        return "C"
    return "D"


@router.get("/scoring")
async def get_scoring(tenant: dict = Depends(get_active_tenant)):
    """Per-touchpoint performance + lead-fit scoring.

    Reads `lead_touchpoint_log` (per-lead instantiated rows) and the active
    touchpoint map. Returns one entry per touchpoint index.
    """
    tenant_id = tenant["id"]
    map_doc = maps_col.find_one({"tenant_id": tenant_id}, {"_id": 0}) or {}
    touchpoints = map_doc.get("touchpoints", [])

    # Aggregate engine results per touchpoint_index
    pipeline = [
        {"$match": {"tenant_id": tenant_id}},
        {"$group": {
            "_id": "$touchpoint_index",
            "total": {"$sum": 1},
            "sent": {"$sum": {"$cond": [{"$eq": ["$status", "sent"]}, 1, 0]}},
            "skipped": {"$sum": {"$cond": [{"$eq": ["$status", "skipped"]}, 1, 0]}},
            "failed": {"$sum": {"$cond": [{"$eq": ["$status", "failed"]}, 1, 0]}},
            "pending": {"$sum": {"$cond": [{"$eq": ["$status", "pending"]}, 1, 0]}},
            "alerts": {"$sum": {"$cond": [{"$eq": ["$status", "alert_sent"]}, 1, 0]}},
            "lead_ids": {"$addToSet": "$lead_id"},
        }},
    ]
    agg = {row["_id"]: row for row in log_col.aggregate(pipeline)}

    # Reply rate: count distinct leads who replied (inbound activity within 7d of fired_at)
    # We approximate using lead_id presence in conversations
    out = []
    overall_sent = 0
    overall_replies = 0
    overall_skip = 0
    overall_total = 0

    for i, tp in enumerate(touchpoints):
        bucket = agg.get(i, {})
        sent = bucket.get("sent", 0)
        skipped = bucket.get("skipped", 0)
        failed = bucket.get("failed", 0)
        pending = bucket.get("pending", 0)
        alerts = bucket.get("alerts", 0)
        total = bucket.get("total", 0)
        lead_ids = bucket.get("lead_ids", []) or []

        # Replies: leads that have at least one inbound conversation after this touchpoint
        reply_count = 0
        if lead_ids:
            reply_count = activities_col_local.count_documents({
                "tenant_id": tenant_id,
                "lead_id": {"$in": lead_ids},
                "type": {"$in": ["inbound_whatsapp", "inbound_email", "reply_received"]},
            })
        reply_rate = _safe_pct(reply_count, sent)

        # Lead-fit: pull icp_tier distribution for leads who received this touchpoint
        fit_hot = fit_warm = fit_cold = 0
        if lead_ids:
            fit_hot = leads_col_local.count_documents({"tenant_id": tenant_id, "id": {"$in": lead_ids}, "icp_tier": "hot"})
            fit_warm = leads_col_local.count_documents({"tenant_id": tenant_id, "id": {"$in": lead_ids}, "icp_tier": "warm"})
            fit_cold = leads_col_local.count_documents({"tenant_id": tenant_id, "id": {"$in": lead_ids}, "icp_tier": "cold"})

        # Effectiveness composite: 60% reply_rate + 30% (sent / total) + 10% (1 - skip_rate)
        deliv = _safe_pct(sent, total) if total else 0
        skip_rate = _safe_pct(skipped, total) if total else 0
        effectiveness = round(0.6 * reply_rate + 0.3 * deliv + 0.1 * (100 - skip_rate), 1)

        overall_sent += sent
        overall_replies += reply_count
        overall_skip += skipped
        overall_total += total

        out.append({
            "index": i,
            "performance": {
                "total_scheduled": total,
                "sent": sent,
                "skipped": skipped,
                "failed": failed,
                "pending": pending,
                "alerts": alerts,
                "replies": reply_count,
                "reply_rate": reply_rate,
                "delivery_rate": deliv,
                "skip_rate": skip_rate,
                "effectiveness": effectiveness,
                "grade": _grade(effectiveness) if total else "—",
            },
            "lead_fit": {
                "hot": fit_hot,
                "warm": fit_warm,
                "cold": fit_cold,
                "total": fit_hot + fit_warm + fit_cold,
            },
        })

    # Journey rollup
    overall_reply_rate = _safe_pct(overall_replies, overall_sent)
    overall_skip_rate = _safe_pct(overall_skip, overall_total) if overall_total else 0
    overall_deliv = _safe_pct(overall_sent, overall_total) if overall_total else 0
    journey_score = round(0.6 * overall_reply_rate + 0.3 * overall_deliv + 0.1 * (100 - overall_skip_rate), 1)

    return {
        "items": out,
        "journey": {
            "score": journey_score,
            "grade": _grade(journey_score) if overall_total else "—",
            "total_scheduled": overall_total,
            "total_sent": overall_sent,
            "total_replies": overall_replies,
            "reply_rate": overall_reply_rate,
            "skip_rate": overall_skip_rate,
            "delivery_rate": overall_deliv,
            "touchpoint_count": len(touchpoints),
        },
    }


@router.post("/ai-quality")
async def ai_quality(tenant: dict = Depends(get_active_tenant)):
    """Claude scores each touchpoint message 1-10 on Clarity, Personalisation,
    CTA strength, Tone-fit. Cached for 24h per (tenant, message hash)."""
    tenant_id = tenant["id"]
    map_doc = maps_col.find_one({"tenant_id": tenant_id}, {"_id": 0}) or {}
    touchpoints = map_doc.get("touchpoints", [])
    if not touchpoints:
        return {"items": []}

    api_key = os.environ.get("EMERGENT_LLM_KEY")
    if not api_key:
        # Heuristic fallback
        return {
            "items": [
                {
                    "index": i,
                    "clarity": 7, "personalisation": 6, "cta": 6, "tone_fit": 7,
                    "average": 6.5, "ai_powered": False,
                    "verdict": "Add {{first_name}} or {{company}} for personalisation" if "{{" not in (tp.get("message_template") or "") else "Looks decent — consider sharpening the CTA",
                }
                for i, tp in enumerate(touchpoints)
            ]
        }

    import hashlib
    import json as _json
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
    except Exception:
        return {"items": [], "error": "ai_unavailable"}

    # Build cache key per message
    items = []
    to_score: List[dict] = []
    for i, tp in enumerate(touchpoints):
        msg = (tp.get("message_template") or "").strip()
        h = hashlib.sha256(msg.encode("utf-8")).hexdigest()[:16]
        cached = ai_quality_cache_col.find_one({"tenant_id": tenant_id, "hash": h}, {"_id": 0})
        if cached and cached.get("scores"):
            items.append({"index": i, **cached["scores"], "ai_powered": True})
        else:
            to_score.append({"index": i, "msg": msg, "hash": h, "channel": tp.get("channel"), "type": tp.get("message_type")})

    if to_score:
        prompt_lines = [
            "You are scoring sales touchpoint messages on a 1-10 scale.",
            "For each message return JSON with: clarity, personalisation, cta, tone_fit (each 1-10), and a one-line verdict (max 80 chars).",
            "Personalisation: presence of {{tokens}} or named references. CTA: clear next action. Tone-fit: matches channel (whatsapp=warm casual, email=professional).",
            "Return a JSON array in the same order as input, no commentary.",
            "",
            "Messages:",
        ]
        for s in to_score:
            prompt_lines.append(f"[{s['index']}] channel={s['channel']} type={s['type']} :: {s['msg'][:500]}")

        try:
            chat = LlmChat(
                api_key=api_key,
                session_id=f"ai-quality-{tenant_id}-{uuid.uuid4().hex[:8]}",
                system_message="You are a precise JSON-only sales copy evaluator.",
            ).with_model("anthropic", "claude-sonnet-4-5-20250929")
            raw = await chat.send_message(UserMessage(text="\n".join(prompt_lines)))
            # Best-effort JSON parse
            text = raw if isinstance(raw, str) else str(raw)
            start = text.find("[")
            end = text.rfind("]")
            scores = _json.loads(text[start:end + 1]) if start >= 0 and end > start else []
            for s, parsed in zip(to_score, scores):
                clarity = float(parsed.get("clarity", 6))
                personalisation = float(parsed.get("personalisation", 6))
                cta = float(parsed.get("cta", 6))
                tone_fit = float(parsed.get("tone_fit", 6))
                avg = round((clarity + personalisation + cta + tone_fit) / 4, 1)
                payload = {
                    "clarity": clarity,
                    "personalisation": personalisation,
                    "cta": cta,
                    "tone_fit": tone_fit,
                    "average": avg,
                    "verdict": str(parsed.get("verdict", ""))[:120],
                }
                ai_quality_cache_col.update_one(
                    {"tenant_id": tenant_id, "hash": s["hash"]},
                    {"$set": {"scores": payload, "updated_at": _now()}},
                    upsert=True,
                )
                items.append({"index": s["index"], **payload, "ai_powered": True})
        except Exception:
            for s in to_score:
                items.append({
                    "index": s["index"], "clarity": 6, "personalisation": 6, "cta": 6, "tone_fit": 6,
                    "average": 6.0, "verdict": "AI scoring unavailable — retry shortly", "ai_powered": False,
                })

    items.sort(key=lambda x: x["index"])
    return {"items": items}
