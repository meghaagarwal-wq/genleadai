"""Sprint 4 / Iter 37 — Touchpoint Mapping completion tests.

Covers:
- 32-touchpoint cap (raised from 30)
- Version history with 5-version retention cap
- /map/versions, /map/versions/{id}/restore (incl. 403 / 404)
- Claude-powered /import-document (DOCX), size/format/role gating
- Audit-log integration for save / import / restore
"""
import io
import os
import time
import pytest
import requests

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL").rstrip("/")
DEMO_TENANT = "ten_demo"
ADMIN = {"email": "admin@demo.com", "password": "Demo1234!"}
SALES = {"email": "sarah@demo.com", "password": "Demo1234!"}


def _login(creds):
    r = requests.post(f"{BASE_URL}/api/auth/login", json=creds, timeout=15)
    assert r.status_code == 200, f"Login failed: {r.status_code} {r.text}"
    return r.json().get("token") or r.json().get("access_token")


@pytest.fixture(scope="module")
def admin_session():
    s = requests.Session()
    tok = _login(ADMIN)
    s.headers.update({
        "Authorization": f"Bearer {tok}",
        "X-Tenant-Id": DEMO_TENANT,
        "Content-Type": "application/json",
    })
    return s


@pytest.fixture(scope="module")
def sales_session():
    s = requests.Session()
    tok = _login(SALES)
    s.headers.update({
        "Authorization": f"Bearer {tok}",
        "X-Tenant-Id": DEMO_TENANT,
        "Content-Type": "application/json",
    })
    return s


def _tp(i, msg="hello"):
    return {
        "index": i,
        "day": float(i),
        "hour": 9,
        "channel": "whatsapp",
        "message_type": "intro" if i == 0 else "follow_up",
        "aria_role": "autonomous",
        "trigger": "",
        "message_template": f"{msg} #{i}",
    }


def _payload(n, marker="default"):
    return {
        "template_id": "tpl_standard",
        "is_customised": True,
        "touchpoints": [_tp(i, marker) for i in range(n)],
    }


# ─── Templates ───────────────────────────────────────────────────────────────
class TestTemplates:
    def test_8_templates_present(self, admin_session):
        r = admin_session.get(f"{BASE_URL}/api/touchpoints/templates")
        assert r.status_code == 200, r.text
        ids = {t["id"] for t in r.json()["templates"]}
        expected = {
            "tpl_fast_conversion", "tpl_warm_nurture", "tpl_considered_sale",
            "tpl_enterprise_track", "tpl_urgency_led", "tpl_high_intent_qualifier",
            "tpl_abandoned_interest", "tpl_standard",
        }
        assert ids == expected, f"diff: {ids ^ expected}"


# ─── 32-touchpoint cap ───────────────────────────────────────────────────────
class TestMaxCap:
    def test_get_map(self, admin_session):
        r = admin_session.get(f"{BASE_URL}/api/touchpoints/map")
        assert r.status_code == 200
        # response shape consistent
        assert "map" in r.json()

    def test_save_32_succeeds(self, admin_session):
        r = admin_session.post(f"{BASE_URL}/api/touchpoints/map", json=_payload(32, "cap32"))
        assert r.status_code == 200, r.text
        m = r.json()["map"]
        assert m["touchpoint_count"] == 32

    def test_save_33_fails_400(self, admin_session):
        r = admin_session.post(f"{BASE_URL}/api/touchpoints/map", json=_payload(33, "cap33"))
        assert r.status_code == 400
        assert "32" in r.text


# ─── Version history & retention ─────────────────────────────────────────────
class TestVersionHistory:
    def test_save_then_versions_endpoint(self, admin_session):
        # Save initial map
        admin_session.post(f"{BASE_URL}/api/touchpoints/map", json=_payload(3, "v_initial"))
        # Save second time → should create one snapshot of v_initial
        admin_session.post(f"{BASE_URL}/api/touchpoints/map", json=_payload(3, "v_second"))
        r = admin_session.get(f"{BASE_URL}/api/touchpoints/map/versions")
        assert r.status_code == 200, r.text
        versions = r.json()["versions"]
        assert isinstance(versions, list)
        assert len(versions) >= 1, "expected at least 1 snapshot after re-save"

    def test_retention_cap_5(self, admin_session):
        # Save 6 distinct maps → versions endpoint must return exactly 5 (oldest pruned)
        for i in range(6):
            r = admin_session.post(
                f"{BASE_URL}/api/touchpoints/map",
                json=_payload(3 + i, f"retain_{i}"),
            )
            assert r.status_code == 200, r.text
            time.sleep(0.15)  # ensure distinct created_at ordering
        r = admin_session.get(f"{BASE_URL}/api/touchpoints/map/versions")
        assert r.status_code == 200
        versions = r.json()["versions"]
        assert len(versions) == 5, f"expected 5 retained, got {len(versions)}"
        # newest first
        ts = [v["created_at"] for v in versions]
        assert ts == sorted(ts, reverse=True), "versions not sorted newest-first"

    def test_restore_version(self, admin_session):
        # Save a uniquely-marked map → save another → restore the snapshot
        admin_session.post(f"{BASE_URL}/api/touchpoints/map", json=_payload(4, "to_be_restored"))
        time.sleep(0.2)
        admin_session.post(f"{BASE_URL}/api/touchpoints/map", json=_payload(5, "new_active"))
        time.sleep(0.2)
        v = admin_session.get(f"{BASE_URL}/api/touchpoints/map/versions").json()["versions"]
        # Find the "to_be_restored" snapshot
        target = None
        for ver in v:
            if any("to_be_restored" in (tp.get("message_template") or "") for tp in ver.get("touchpoints", [])):
                target = ver
                break
        assert target is not None, "could not find target version to restore"
        r = admin_session.post(
            f"{BASE_URL}/api/touchpoints/map/versions/{target['id']}/restore"
        )
        assert r.status_code == 200, r.text
        m = r.json()["map"]
        assert m["touchpoint_count"] == 4
        assert m.get("restored_from") == target["id"]
        # Verify GET reflects the restored state
        g = admin_session.get(f"{BASE_URL}/api/touchpoints/map").json()["map"]
        assert g["touchpoint_count"] == 4
        assert any("to_be_restored" in tp["message_template"] for tp in g["touchpoints"])

    def test_restore_404_for_unknown(self, admin_session):
        r = admin_session.post(
            f"{BASE_URL}/api/touchpoints/map/versions/ver_does_not_exist/restore"
        )
        assert r.status_code == 404

    def test_restore_403_for_sales_rep(self, sales_session):
        r = sales_session.post(
            f"{BASE_URL}/api/touchpoints/map/versions/ver_xyz/restore"
        )
        assert r.status_code == 403


# ─── Document upload (Claude) ────────────────────────────────────────────────
def _make_docx_bytes(steps=5):
    """Generate a small DOCX with N numbered steps describing a sequence."""
    from docx import Document
    d = Document()
    d.add_heading("Outreach Sequence", level=1)
    canned = [
        ("Step 1 - Day 0", "WhatsApp intro: Hi {{lead_name}}, thanks for your interest in {{product}}."),
        ("Step 2 - Day 1", "Follow-up email: Sharing more details about pricing and timelines."),
        ("Step 3 - Day 3", "Call reminder for sales rep to phone the lead manually."),
        ("Step 4 - Day 5", "LinkedIn nudge with social proof - link to case study."),
        ("Step 5 - Day 7", "Final WhatsApp meeting CTA - book a 30-min demo."),
        ("Step 6 - Day 10", "Re-engagement message if no reply."),
        ("Step 7 - Day 14", "Closure - thanks and goodbye."),
    ]
    for i in range(min(steps, len(canned))):
        d.add_paragraph(canned[i][0], style="Heading 2")
        d.add_paragraph(canned[i][1])
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


@pytest.fixture(scope="module")
def docx_5_bytes():
    return _make_docx_bytes(5)


class TestImportDocument:
    def _post_file(self, session, filename, content, content_type):
        files = {"file": (filename, content, content_type)}
        h = {k: v for k, v in session.headers.items() if k.lower() != "content-type"}
        return requests.post(
            f"{BASE_URL}/api/touchpoints/import-document",
            files=files, headers=h, timeout=60,
        )

    def test_docx_5_steps_returns_preview(self, admin_session, docx_5_bytes):
        r = self._post_file(
            admin_session, "test_sequence.docx", docx_5_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert r.status_code == 200, f"status={r.status_code} body={r.text[:600]}"
        body = r.json()
        for k in ("preview", "extracted_count", "truncated", "source_filename"):
            assert k in body, f"missing key {k} in {body}"
        assert body["source_filename"] == "test_sequence.docx"
        assert body["truncated"] is False
        assert isinstance(body["preview"], list)
        # Claude may extract anywhere from 3-7 — be lenient but require at least 3
        assert len(body["preview"]) >= 3, f"too few extracted: {body['extracted_count']}"
        # Normalisation: channel/role/type are within allowed values
        allowed_ch = {"whatsapp", "email", "call_reminder", "linkedin_nudge"}
        allowed_role = {"autonomous", "alert_human"}
        for tp in body["preview"]:
            assert tp["channel"] in allowed_ch, f"bad channel {tp['channel']}"
            assert tp["aria_role"] in allowed_role, f"bad role {tp['aria_role']}"
            assert tp["message_template"], "message_template should default"

    def test_empty_file_400(self, admin_session):
        r = self._post_file(admin_session, "empty.docx", b"",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert r.status_code == 400

    def test_unsupported_extension_400(self, admin_session):
        r = self._post_file(admin_session, "notes.txt",
                            b"Step 1 - Day 0 - whatsapp intro " * 5, "text/plain")
        assert r.status_code == 400

    def test_over_10mb_400(self, admin_session):
        big = b"A" * (10 * 1024 * 1024 + 10)
        r = self._post_file(admin_session, "huge.pdf", big, "application/pdf")
        assert r.status_code == 400
        assert "10MB" in r.text or "limit" in r.text.lower()

    def test_sales_rep_403(self, sales_session, docx_5_bytes):
        r = self._post_file(sales_session, "blocked.docx", docx_5_bytes,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert r.status_code == 403


# ─── Audit log integration ───────────────────────────────────────────────────
class TestAuditIntegration:
    def test_save_writes_audit(self, admin_session):
        admin_session.post(f"{BASE_URL}/api/touchpoints/map", json=_payload(3, "audit_save"))
        time.sleep(0.5)
        r = admin_session.get(
            f"{BASE_URL}/api/audit-log",
            params={"action": "touchpoint_map.updated"},
        )
        assert r.status_code == 200, r.text
        body = r.json()
        assert body["total"] >= 1
        assert any(e["action"] == "touchpoint_map.updated" for e in body["entries"])

    def test_import_writes_audit(self, admin_session, docx_5_bytes):
        files = {"file": ("audit_check.docx", docx_5_bytes,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        h = {k: v for k, v in admin_session.headers.items() if k.lower() != "content-type"}
        ir = requests.post(f"{BASE_URL}/api/touchpoints/import-document",
                           files=files, headers=h, timeout=60)
        if ir.status_code != 200:
            pytest.skip(f"import failed upstream: {ir.status_code}")
        time.sleep(0.5)
        r = admin_session.get(
            f"{BASE_URL}/api/audit-log",
            params={"action": "touchpoint_map.imported_from_document"},
        )
        assert r.status_code == 200
        assert r.json()["total"] >= 1

    def test_restore_writes_audit(self, admin_session):
        admin_session.post(f"{BASE_URL}/api/touchpoints/map", json=_payload(3, "audit_restore_a"))
        time.sleep(0.2)
        admin_session.post(f"{BASE_URL}/api/touchpoints/map", json=_payload(4, "audit_restore_b"))
        v = admin_session.get(f"{BASE_URL}/api/touchpoints/map/versions").json()["versions"]
        if not v:
            pytest.skip("no versions available for restore audit test")
        admin_session.post(f"{BASE_URL}/api/touchpoints/map/versions/{v[0]['id']}/restore")
        time.sleep(0.5)
        r = admin_session.get(
            f"{BASE_URL}/api/audit-log",
            params={"action": "touchpoint_map.restored"},
        )
        assert r.status_code == 200
        assert r.json()["total"] >= 1
